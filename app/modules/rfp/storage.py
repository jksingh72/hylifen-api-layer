"""
File storage, text extraction, and .docx generation helpers for the RFP/RFI module.
Kept self-contained (no imports from other modules) so this module stays isolated.
"""
import os
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

RFP_UPLOAD_DIR = Path("rfp_uploads")
RFP_GENERATED_DIR = Path("rfp_generated")
RFP_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
RFP_GENERATED_DIR.mkdir(parents=True, exist_ok=True)


async def save_rfp_upload(rfp_id: str, upload_file: UploadFile) -> str:
    """Saves an uploaded source document under rfp_uploads/{rfp_id}/ and returns its path."""
    rfp_dir = RFP_UPLOAD_DIR / rfp_id
    rfp_dir.mkdir(parents=True, exist_ok=True)
    file_path = rfp_dir / upload_file.filename
    try:
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
        return str(file_path.absolute())
    finally:
        await upload_file.close()


def delete_rfp_file(storage_path: str) -> None:
    if os.path.exists(storage_path):
        os.remove(storage_path)


def extract_text_from_file(file_path: str) -> str:
    """Extracts plain text from a source document (pdf/docx/xlsx/txt) for agent analysis."""
    if not os.path.exists(file_path):
        return f"Error: File not found at {file_path}"

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if ext in (".docx", ".doc"):
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(file_path)
            return df.to_string()
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as exc:
        return f"Error parsing {ext} file: {exc}"


def generate_format_docx(rfp_id: str, version: int, section_titles: list[str]) -> str:
    """Creates a .docx containing only section headings (no body content yet)."""
    rfp_dir = RFP_GENERATED_DIR / rfp_id
    rfp_dir.mkdir(parents=True, exist_ok=True)
    file_path = rfp_dir / f"format_v{version}.docx"

    doc = Document()
    doc.add_heading("Response Format", level=1)
    for title in section_titles:
        doc.add_heading(title, level=2)
        doc.add_paragraph("")
    doc.save(file_path)
    return str(file_path.absolute())


def read_format_sections(file_path: str) -> list[str]:
    """Reads section (level-2 heading) titles back out of a format .docx."""
    doc = Document(file_path)
    return [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]


def generate_response_docx(rfp_id: str, version: int, sections: dict[str, str]) -> str:
    """Creates a .docx with each section heading followed by its written content."""
    rfp_dir = RFP_GENERATED_DIR / rfp_id
    rfp_dir.mkdir(parents=True, exist_ok=True)
    file_path = rfp_dir / f"response_v{version}.docx"

    doc = Document()
    doc.add_heading("Response Document", level=1)
    for title, content in sections.items():
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    doc.save(file_path)
    return str(file_path.absolute())
